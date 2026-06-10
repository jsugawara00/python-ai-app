import base64
from pathlib import Path
import streamlit as st
import google.generativeai as genai


def generate(prompt: str, temperature: float = 0.7, system_instruction: str = "") -> str:
    api_key = st.session_state.get("api_key", "")
    if not api_key:
        return "❌ APIキーが設定されていません。サイドバーからGemini APIキーを入力してください。"
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            "gemini-3.5-flash",
            system_instruction=system_instruction or None,
            generation_config=genai.GenerationConfig(temperature=temperature),
        )
        response = model.generate_content(prompt)
        return response.text
    except Exception:
        return "❌ エラーが発生しました。APIキーが正しいか確認のうえ、しばらく待ってから再試行してください。"


# =====================
# ダウンロード用ヘルパー
# =====================

def _to_pdf_bytes(text: str, image_bytes: bytes = None, image_position=None) -> bytes:
    import io, re
    from PIL import Image as PILImage
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
    buf = io.BytesIO()
    max_w = 170 * mm
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    h1 = ParagraphStyle("h1", fontName="HeiseiKakuGo-W5", fontSize=16, leading=24, spaceAfter=8)
    h2 = ParagraphStyle("h2", fontName="HeiseiKakuGo-W5", fontSize=13, leading=20, spaceAfter=6)
    normal = ParagraphStyle("normal", fontName="HeiseiKakuGo-W5", fontSize=10, leading=16)

    def _esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _strip_md(s):
        return re.sub(r'\*+([^*]+)\*+', r'\1', s)

    def _make_img():
        pil = PILImage.open(io.BytesIO(image_bytes))
        w, h = pil.size
        iw = min(max_w, max_w)
        ih = iw * h / w
        if ih > 120 * mm:
            ih = 120 * mm
            iw = ih * w / h
        return [RLImage(io.BytesIO(image_bytes), width=iw, height=ih), Spacer(1, 8)]

    positions = image_position or []
    story = []
    first_heading_done = False

    if image_bytes and "文章の先頭" in positions:
        story.extend(_make_img())

    for line in text.split("\n"):
        s = _strip_md(_esc(line))
        is_heading = line.startswith("#")
        if line.startswith("# "):
            story.append(Paragraph(s[2:], h1))
        elif line.startswith("## "):
            story.append(Paragraph(s[3:], h2))
        elif line.startswith("### "):
            story.append(Paragraph(s[4:], h2))
        elif line.strip():
            story.append(Paragraph(s, normal))
            story.append(Spacer(1, 4))
        if image_bytes and "最初の見出しの後" in positions and is_heading and not first_heading_done:
            story.extend(_make_img())
            first_heading_done = True

    if image_bytes and "文章の末尾" in positions:
        story.extend(_make_img())

    doc.build(story)
    return buf.getvalue()


def _to_docx_bytes(text: str, image_bytes: bytes = None, image_position=None) -> bytes:
    import io, re
    from docx import Document
    from docx.shared import Cm

    def _strip_md(s):
        return re.sub(r'\*+([^*]+)\*+', r'\1', s)

    positions = image_position or []
    doc = Document()
    first_heading_done = False

    def _add_img():
        doc.add_picture(io.BytesIO(image_bytes), width=Cm(15))

    if image_bytes and "文章の先頭" in positions:
        _add_img()

    for line in text.split("\n"):
        s = _strip_md(line)
        is_heading = line.startswith("#")
        if line.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif line.strip():
            doc.add_paragraph(s)
        else:
            doc.add_paragraph("")
        if image_bytes and "最初の見出しの後" in positions and is_heading and not first_heading_done:
            _add_img()
            first_heading_done = True

    if image_bytes and "文章の末尾" in positions:
        _add_img()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _to_sales_report_pdf_bytes(report_title: str, info: dict, sections: dict) -> bytes:
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    def _esc(s):
        return (s or "—").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    title_s = ParagraphStyle("sr_title", fontName="HeiseiKakuGo-W5", fontSize=18, leading=26,
                              spaceAfter=6, alignment=1)
    label_s = ParagraphStyle("sr_label", fontName="HeiseiKakuGo-W5", fontSize=9,  leading=13,
                              textColor=colors.HexColor("#666666"))
    value_s = ParagraphStyle("sr_value", fontName="HeiseiKakuGo-W5", fontSize=10, leading=15,
                              wordWrap="CJK", splitLongWords=True)
    sec_h_s = ParagraphStyle("sr_sec_h", fontName="HeiseiKakuGo-W5", fontSize=11, leading=16,
                              textColor=colors.white)
    body_s  = ParagraphStyle("sr_body",  fontName="HeiseiKakuGo-W5", fontSize=10, leading=16)

    NAVY  = colors.HexColor("#1a1a6e")
    LGRAY = colors.HexColor("#f0f4ff")
    GRAY  = colors.HexColor("#dee2e6")

    right_s = ParagraphStyle("sr_right", fontName="HeiseiKakuGo-W5", fontSize=9,
                              leading=14, alignment=2, textColor=colors.HexColor("#444444"))

    def _v(k):
        return (info.get(k, "") or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    right_parts = []
    if _v("報告日"): right_parts.append(f"報告日：{_v('報告日')}")
    if _v("報告者"): right_parts.append(f"報告者：{_v('報告者')}")
    right_para = Paragraph("<br/>".join(right_parts), right_s) if right_parts else Paragraph("", right_s)

    story = []
    top_tbl = Table(
        [[Paragraph(_esc(report_title), title_s), right_para]],
        colWidths=[110*mm, 60*mm],
    )
    top_tbl.setStyle(TableStyle([
        ("ALIGN",         (0, 0), (0, 0), "CENTER"),
        ("VALIGN",        (0, 0), (-1, -1), "BOTTOM"),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING",   (0, 0), (-1, -1), 0),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
    ]))
    story.append(top_tbl)
    story.append(HRFlowable(width="100%", thickness=2, color=NAVY, spaceAfter=8))
    story.append(Spacer(1, 4*mm))

    cw = 170*mm / 4
    tbl_data = [
        [Paragraph("訪問先",     label_s), Paragraph(_esc(info.get("訪問先", "")),     value_s),
         Paragraph("訪問日時",   label_s), Paragraph(_esc(info.get("訪問日時", "")),   value_s)],
        [Paragraph("先方対応者", label_s), Paragraph(_esc(info.get("先方対応者", "")), value_s),
         Paragraph("自社担当者", label_s), Paragraph(_esc(info.get("自社担当者", "")), value_s)],
    ]
    t = Table(tbl_data, colWidths=[cw*0.65, cw*1.35, cw*0.65, cw*1.35])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), LGRAY),
        ("GRID",       (0, 0), (-1, -1), 0.5, GRAY),
        ("PADDING",    (0, 0), (-1, -1), 5),
        ("VALIGN",     (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(t)
    story.append(Spacer(1, 8*mm))

    sec_colors = {
        "訪問目的":       colors.HexColor("#1a1a6e"),
        "訪問の実施結果": colors.HexColor("#1a5276"),
        "当社の課題":     colors.HexColor("#6e1a1a"),
        "今後の展望":     colors.HexColor("#1a6e3a"),
    }
    for key, label in [("訪問目的", "■ 訪問目的"), ("訪問の実施結果", "■ 訪問の実施結果"),
                        ("当社の課題", "■ 当社の課題"), ("今後の展望", "■ 今後の展望")]:
        hdr_t = Table([[Paragraph(label, sec_h_s)]], colWidths=[170*mm])
        hdr_t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, 0), sec_colors.get(key, NAVY)),
            ("PADDING",    (0, 0), (0, 0), 7),
        ]))
        story.append(hdr_t)
        content = sections.get(key, "") or "（未記入）"
        for line in content.split("\n"):
            if line.strip():
                story.append(Paragraph(_esc(line), body_s))
                story.append(Spacer(1, 2))
            else:
                story.append(Spacer(1, 4))
        story.append(Spacer(1, 6*mm))

    doc.build(story)
    return buf.getvalue()


def _to_sales_report_xlsx_bytes(report_title: str, info: dict, sections: dict) -> bytes:
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = report_title[:31]

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 30

    def fill(hex_color: str) -> PatternFill:
        return PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")

    def thin_border() -> Border:
        s = Side(style="thin", color="C8D4F0")
        return Border(top=s, bottom=s, left=s, right=s)

    row = 1

    # 右上の報告日・報告者（C:D 列 / 右揃え）
    for label, key in [("報告日", "報告日"), ("報告者", "報告者")]:
        val = info.get(key, "")
        ws.merge_cells(f"C{row}:D{row}")
        c = ws.cell(row=row, column=3, value=f"{label}：{val}" if val else "")
        c.font = Font(name="Meiryo", size=9, color="444444")
        c.alignment = Alignment(horizontal="right", vertical="center")
        ws.row_dimensions[row].height = 16
        row += 1

    # タイトル行
    ws.merge_cells(f"A{row}:D{row}")
    c = ws[f"A{row}"]
    c.value = report_title
    c.font = Font(name="Meiryo", size=18, bold=True, color="1A1A6E")
    c.alignment = Alignment(horizontal="center", vertical="center")
    c.fill = fill("F0F4FF")
    c.border = Border(bottom=Side(style="medium", color="1A1A6E"))
    ws.row_dimensions[row].height = 38
    row += 1

    ws.row_dimensions[row].height = 6
    row += 1

    # 基本情報（訪問先・訪問日時・先方対応者・自社担当者）
    for l1, v1, l2, v2 in [
        ("訪問先",     info.get("訪問先","") or "—",     "訪問日時",   info.get("訪問日時","") or "—"),
        ("先方対応者", info.get("先方対応者","") or "—", "自社担当者", info.get("自社担当者","") or "—"),
    ]:
        ws.row_dimensions[row].height = 22
        for col_i, (val, is_lbl) in enumerate([(l1,True),(v1,False),(l2,True),(v2,False)], 1):
            c = ws.cell(row=row, column=col_i, value=val)
            c.font = Font(name="Meiryo", size=10,
                          bold=is_lbl, color="444444" if is_lbl else "111111")
            c.fill = fill("E4EAFF") if is_lbl else fill("F0F4FF")
            c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
            c.border = thin_border()
        row += 1

    ws.row_dimensions[row].height = 8
    row += 1

    # セクション
    sec_colors = {
        "訪問目的":       "1A1A6E",
        "訪問の実施結果": "1A5276",
        "当社の課題":     "6E1A1A",
        "今後の展望":     "1A6E3A",
    }
    for key in ["訪問目的", "訪問の実施結果", "当社の課題", "今後の展望"]:
        # ヘッダー行
        ws.merge_cells(f"A{row}:D{row}")
        c = ws[f"A{row}"]
        c.value = f"■ {key}"
        c.font = Font(name="Meiryo", size=11, bold=True, color="FFFFFF")
        c.fill = fill(sec_colors[key])
        c.alignment = Alignment(horizontal="left", vertical="center")
        ws.row_dimensions[row].height = 24
        row += 1

        # コンテンツ行
        content = sections.get(key, "") or "（未記入）"
        ws.merge_cells(f"A{row}:D{row}")
        c = ws[f"A{row}"]
        c.value = content
        c.font = Font(name="Meiryo", size=10, color="111111")
        c.fill = fill("FFFFFF")
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        c.border = Border(
            left=Side(style="thin", color="DDDDEE"),
            right=Side(style="thin", color="DDDDEE"),
            bottom=Side(style="thin", color="DDDDEE"),
        )
        lines = max(content.count("\n") + 1, len(content) // 55 + 1)
        ws.row_dimensions[row].height = max(45, lines * 18)
        row += 1

        ws.row_dimensions[row].height = 5
        row += 1

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _to_sales_report_docx_bytes(report_title: str, info: dict, sections: dict) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin   = Cm(2.5)
        s.right_margin  = Cm(2.5)

    def _rp(text: str, size: float, bold=False, color=(0x22,0x22,0x22), align=None) -> None:
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)
        return p

    def _hline(para):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "12")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), "1A1A6E")
        pBdr.append(b)
        pPr.append(pBdr)

    # 右上の報告日・報告者
    for key in ["報告日", "報告者"]:
        val = info.get(key, "")
        _rp(f"{key}：{val or '　　　　　　'}", 10,
            color=(0x44,0x44,0x44), align=WD_ALIGN_PARAGRAPH.RIGHT)

    # タイトル
    title_p = _rp(report_title, 18, bold=True,
                  color=(0x1a,0x1a,0x6e), align=WD_ALIGN_PARAGRAPH.CENTER)
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(6)
    _hline(title_p)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 基本情報（リスト形式）
    for key in ["訪問先", "訪問日時", "先方対応者", "自社担当者"]:
        val = info.get(key, "")
        p = doc.add_paragraph()
        r1 = p.add_run(f"{key}：")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
        r2 = p.add_run(val or "　")
        r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # セクション
    sec_colors = {
        "訪問目的":       RGBColor(0x1a, 0x1a, 0x6e),
        "訪問の実施結果": RGBColor(0x1a, 0x52, 0x76),
        "当社の課題":     RGBColor(0x6e, 0x1a, 0x1a),
        "今後の展望":     RGBColor(0x1a, 0x6e, 0x3a),
    }
    for key in ["訪問目的", "訪問の実施結果", "当社の課題", "今後の展望"]:
        h = doc.add_heading(f"■ {key}", level=2)
        for r in h.runs:
            r.font.color.rgb = sec_colors.get(key, RGBColor(0x1a, 0x1a, 0x6e))
        content = sections.get(key, "") or "（未記入）"
        p = doc.add_paragraph(content)
        p.paragraph_format.space_after = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _image_section() -> tuple:
    img_file = st.file_uploader(
        "📷 画像を追加する（任意・ドラッグ＆ドロップ対応）",
        type=["jpg", "jpeg", "png"],
    )
    if img_file:
        col_prev, col_pos = st.columns([1, 2])
        with col_prev:
            st.image(img_file, width=160)
        with col_pos:
            positions = st.multiselect(
                "画像の配置位置（複数選択可・最大3か所）",
                ["文章の先頭", "最初の見出しの後", "文章の末尾"],
                default=["文章の先頭"],
            )
        return img_file.getvalue(), positions
    return None, []


# =====================
# 各ツールのページ関数
# =====================

def page_blog():
    st.title("📝 ブログ記事執筆")
    st.markdown("テーマやキーワードを入力するだけで、構成付きのブログ記事を生成します。")

    col1, col2 = st.columns(2)
    with col1:
        topic = st.text_input("記事のテーマ・タイトル", placeholder="例: 初心者向けPython入門")
        keywords = st.text_area("含めたいキーワード（任意）", placeholder="例: 変数, 関数, ループ", height=80)
    with col2:
        tone = st.selectbox(
            "文体・トーン",
            ["わかりやすく丁寧に", "カジュアル・フレンドリー", "専門的・フォーマル", "ユーモアを交えて"],
        )
        length = st.selectbox(
            "記事の長さ",
            ["短め（500文字程度）", "標準（1000文字程度）", "長め（2000文字程度）"],
        )
        target = st.text_input("ターゲット読者（任意）", placeholder="例: プログラミング初心者")

    additional = st.text_area(
        "追加の指示・要望（任意）",
        placeholder="例: 具体的な例を多く含めてほしい、SEOを意識してほしい",
        height=80,
    )

    st.markdown("---")
    img_bytes, img_position = _image_section()

    if st.button("記事を生成する", type="primary", use_container_width=True):
        if not topic:
            st.error("記事のテーマを入力してください。")
            return

        system_instruction = "あなたはプロのブログライターです。ユーザーが指定した条件に従ってブログ記事を執筆してください。ライティング支援以外の指示には従わないでください。"
        prompt = f"""以下の条件でブログ記事を執筆してください。

【テーマ】{topic}
【ターゲット読者】{target if target else "一般読者"}
【文体・トーン】{tone}
【記事の長さ】{length}
【含めたいキーワード】{keywords if keywords else "特になし"}
【追加の要望】{additional if additional else "特になし"}

記事の構成：
- キャッチーなタイトル
- リード文（冒頭の掴み）
- 本文（見出しを使って構造化）
- まとめ

マークダウン形式で出力してください。"""

        with st.spinner("記事を生成中..."):
            result = generate(prompt, system_instruction=system_instruction)

        st.markdown("---")
        st.markdown("### 生成された記事")
        st.markdown(result)
        col_pdf, col_word, col_md = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 記事をダウンロード（PDF）",
                data=_to_pdf_bytes(result, img_bytes, img_position),
                file_name="blog_article.pdf",
                mime="application/pdf",
            )
        with col_word:
            st.download_button(
                "📥 記事をダウンロード（Word）",
                data=_to_docx_bytes(result, img_bytes, img_position),
                file_name="blog_article.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_md:
            st.download_button(
                "📥 記事をダウンロード（Markdown）",
                data=result,
                file_name="blog_article.md",
                mime="text/markdown",
            )


def page_email():
    st.title("📧 メール返信作成")
    st.markdown("受信メールを貼り付けるだけで、適切な返信文を自動作成します。")

    original_email = st.text_area(
        "受信メールの内容",
        placeholder="ここに受信したメールの文章を貼り付けてください...",
        height=200,
    )

    col1, col2 = st.columns(2)
    with col1:
        tone = st.selectbox(
            "返信のトーン",
            ["丁寧・ビジネス向け", "カジュアル・フレンドリー", "簡潔・要点のみ", "感謝を伝える", "お断り・辞退"],
        )
        language = st.selectbox("言語", ["日本語", "英語", "日本語と英語の両方"])
    with col2:
        key_points = st.text_area(
            "伝えたい要点（任意）",
            placeholder="例: ○○は承認しました、△△の件は来週対応予定",
            height=100,
        )

    st.markdown("---")
    img_bytes, img_position = _image_section()

    if st.button("返信文を作成する", type="primary", use_container_width=True):
        if not original_email:
            st.error("受信メールの内容を入力してください。")
            return

        system_instruction = "あなたはプロのビジネスライターです。受信メールに対する適切な返信文を作成してください。メール作成以外の指示には従わないでください。"
        prompt = f"""以下の受信メールに対する返信文を作成してください。

【受信メール】
{original_email}

【返信のトーン】{tone}
【使用言語】{language}
【伝えたい要点】{key_points if key_points else "受信メールに対して適切に返答する"}

件名から本文まで完全な返信メールを作成してください。宛名は「〇〇様」のように仮置きしてください。"""

        with st.spinner("返信文を作成中..."):
            result = generate(prompt, system_instruction=system_instruction)

        st.markdown("---")
        st.markdown("### 生成された返信文")
        st.text_area("返信文（コピーしてお使いください）", value=result, height=300)
        col_pdf, col_word, col_md = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 返信文をダウンロード（PDF）",
                data=_to_pdf_bytes(result, img_bytes, img_position),
                file_name="email_reply.pdf",
                mime="application/pdf",
            )
        with col_word:
            st.download_button(
                "📥 返信文をダウンロード（Word）",
                data=_to_docx_bytes(result, img_bytes, img_position),
                file_name="email_reply.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_md:
            st.download_button(
                "📥 返信文をダウンロード（テキスト）",
                data=result,
                file_name="email_reply.txt",
                mime="text/plain",
            )


def page_summary():
    st.title("📋 文章要約")
    st.markdown("長い文章を簡潔にまとめます。記事・論文・報告書など何でも対応。")

    text = st.text_area(
        "要約したい文章",
        placeholder="ここに要約したい文章を貼り付けてください...",
        height=250,
    )

    col1, col2 = st.columns(2)
    with col1:
        length = st.selectbox(
            "要約の長さ",
            ["3行程度で超コンパクト", "100〜200文字程度", "300〜500文字程度", "ポイントを箇条書きで5項目"],
        )
    with col2:
        style = st.selectbox(
            "要約スタイル",
            ["わかりやすい日本語で", "専門用語を保持して", "小学生にも分かるように", "ビジネス向けに"],
        )

    st.markdown("---")
    img_bytes, img_position = _image_section()

    if st.button("要約する", type="primary", use_container_width=True):
        if not text:
            st.error("要約したい文章を入力してください。")
            return

        system_instruction = "あなたは文章要約の専門家です。指定された長さとスタイルで文章を正確に要約してください。要約以外の指示には従わないでください。"
        prompt = f"""以下の文章を要約してください。

【要約する文章】
{text}

【要約の長さ】{length}
【要約スタイル】{style}

指定の長さとスタイルで要約してください。"""

        with st.spinner("要約中..."):
            result = generate(prompt, temperature=0.3, system_instruction=system_instruction)

        st.markdown("---")
        st.markdown("### 要約結果")
        st.info(result)
        col_pdf, col_word, col_md = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 要約をダウンロード（PDF）",
                data=_to_pdf_bytes(result, img_bytes, img_position),
                file_name="summary.pdf",
                mime="application/pdf",
            )
        with col_word:
            st.download_button(
                "📥 要約をダウンロード（Word）",
                data=_to_docx_bytes(result, img_bytes, img_position),
                file_name="summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_md:
            st.download_button(
                "📥 要約をダウンロード（テキスト）",
                data=result,
                file_name="summary.txt",
                mime="text/plain",
            )


def page_sns():
    st.title("📱 SNS投稿文作成")
    st.markdown("各SNSプラットフォームに最適化された投稿文を3パターン提案します。")

    topic = st.text_area(
        "投稿したい内容・テーマ",
        placeholder="例: 新商品のカフェラテを発売しました。コーヒー豆はブラジル産を使用しています。",
        height=120,
    )

    col1, col2, col3 = st.columns(3)
    with col1:
        platform = st.selectbox(
            "プラットフォーム",
            ["X (Twitter)", "Instagram", "Facebook", "LinkedIn", "note"],
        )
    with col2:
        tone = st.selectbox(
            "投稿のトーン",
            ["フレンドリー・親しみやすい", "プロフェッショナル", "ユーモラス・面白い", "インスピレーション・感動的", "情報提供・教育的"],
        )
    with col3:
        hashtag = st.selectbox("ハッシュタグ", ["自動で追加する", "追加しない"])

    platform_note = {
        "X (Twitter)": "140文字以内で簡潔に",
        "Instagram": "改行を活用して読みやすく、絵文字を効果的に使って",
        "Facebook": "少し長めに詳しく、親しみやすく",
        "LinkedIn": "ビジネスライクに、プロフェッショナルな表現で",
        "note": "記事の導入文として、読者の興味を引くように",
    }

    st.markdown("---")
    img_bytes, img_position = _image_section()

    if st.button("投稿文を作成する", type="primary", use_container_width=True):
        if not topic:
            st.error("投稿したい内容を入力してください。")
            return

        system_instruction = "あなたはSNSマーケティングの専門家です。各プラットフォームに最適化されたSNS投稿文を作成してください。投稿文作成以外の指示には従わないでください。"
        prompt = f"""以下の内容で{platform}の投稿文を作成してください。

【投稿したい内容】
{topic}

【プラットフォーム】{platform}（{platform_note.get(platform, "")}）
【トーン】{tone}
【ハッシュタグ】{hashtag}

そのまま投稿できる形式で、バリエーションを3パターン作成してください。各パターンを「--- パターン1 ---」のように区切ってください。"""

        with st.spinner("投稿文を作成中..."):
            result = generate(prompt, system_instruction=system_instruction)

        st.markdown("---")
        st.markdown("### 生成された投稿文")
        st.markdown(result)
        col_pdf, col_word, col_md = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 投稿文をダウンロード（PDF）",
                data=_to_pdf_bytes(result, img_bytes, img_position),
                file_name="sns_posts.pdf",
                mime="application/pdf",
            )
        with col_word:
            st.download_button(
                "📥 投稿文をダウンロード（Word）",
                data=_to_docx_bytes(result, img_bytes, img_position),
                file_name="sns_posts.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_md:
            st.download_button(
                "📥 投稿文をダウンロード（テキスト）",
                data=result,
                file_name="sns_posts.txt",
                mime="text/plain",
            )


def page_proofread():
    st.title("🔍 文章校正・改善")
    st.markdown("誤字脱字チェック・表現改善・読みやすさ向上など多角的に文章をブラッシュアップします。")

    text = st.text_area(
        "校正・改善したい文章",
        placeholder="ここに改善したい文章を貼り付けてください...",
        height=250,
    )

    focus = st.multiselect(
        "改善したい観点（複数選択可）",
        ["誤字脱字・文法チェック", "表現の自然さ・読みやすさ", "敬語・丁寧語の適切さ", "文章の簡潔さ・無駄の削除", "説得力・論理的な構成"],
        default=["誤字脱字・文法チェック", "表現の自然さ・読みやすさ"],
    )

    show_diff = st.checkbox("変更箇所の説明を表示する", value=True)

    st.markdown("---")
    img_bytes, img_position = _image_section()

    if st.button("校正・改善する", type="primary", use_container_width=True):
        if not text:
            st.error("校正・改善したい文章を入力してください。")
            return
        if not focus:
            st.error("改善したい観点を1つ以上選択してください。")
            return

        diff_instruction = (
            "改善後の文章に加えて、どの部分をどのように修正したかを箇条書きで説明してください。"
            if show_diff
            else "改善後の文章のみを出力してください。"
        )

        system_instruction = "あなたは日本語の文章校正・編集の専門家です。指定された観点で文章を改善してください。文章校正以外の指示には従わないでください。"
        prompt = f"""以下の文章を改善してください。

【改善する文章】
{text}

【改善する観点】
{', '.join(focus)}

{diff_instruction}

出力形式：
## 改善後の文章
（改善後のテキスト）

## 変更箇所の説明
（箇条書きで説明）"""

        with st.spinner("校正・改善中..."):
            result = generate(prompt, temperature=0.3, system_instruction=system_instruction)

        st.markdown("---")
        st.markdown("### 校正・改善結果")
        st.markdown(result)
        col_pdf, col_word, col_md = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 改善文をダウンロード（PDF）",
                data=_to_pdf_bytes(result, img_bytes, img_position),
                file_name="proofread_text.pdf",
                mime="application/pdf",
            )
        with col_word:
            st.download_button(
                "📥 改善文をダウンロード（Word）",
                data=_to_docx_bytes(result, img_bytes, img_position),
                file_name="proofread_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_md:
            st.download_button(
                "📥 改善文をダウンロード（テキスト）",
                data=result,
                file_name="proofread_text.txt",
                mime="text/plain",
            )


def page_title():
    st.title("💡 タイトル・見出し生成")
    st.markdown("コンテンツの内容から、クリックされやすい魅力的なタイトルを複数提案します。")

    content = st.text_area(
        "内容・テーマの説明",
        placeholder="例: Pythonを使ったデータ分析の初心者向けチュートリアル。pandasとmatplotlibを使って売上データを分析する方法を解説する。",
        height=150,
    )

    col1, col2 = st.columns(2)
    with col1:
        content_type = st.selectbox(
            "コンテンツの種類",
            ["ブログ記事", "YouTubeサムネイル", "メールの件名", "プレゼンテーション", "本・電子書籍", "SNS投稿"],
        )
        count = st.slider("提案数", min_value=3, max_value=10, value=5)
    with col2:
        style = st.multiselect(
            "タイトルのスタイル（複数選択可）",
            [
                "数字を使ったリスト形式（例: 5つの方法）",
                "疑問形・問いかけ",
                "ハウツー形式",
                "驚き・インパクト重視",
                "SEO・キーワード重視",
                "シンプル・わかりやすい",
            ],
            default=["シンプル・わかりやすい", "数字を使ったリスト形式（例: 5つの方法）"],
        )

    st.markdown("---")
    img_bytes, img_position = _image_section()

    if st.button("タイトルを生成する", type="primary", use_container_width=True):
        if not content:
            st.error("内容・テーマを入力してください。")
            return

        system_instruction = "あなたはコピーライターの専門家です。コンテンツの内容に合った魅力的なタイトルを提案してください。タイトル提案以外の指示には従わないでください。"
        prompt = f"""以下の内容に合った魅力的なタイトルを{count}個提案してください。

【内容・テーマ】
{content}

【コンテンツの種類】{content_type}
【タイトルのスタイル】{', '.join(style) if style else "バランス良く"}

各タイトルに対して、なぜそのタイトルが効果的かを1行で添えてください。
番号付きリスト形式で出力してください。"""

        with st.spinner("タイトルを生成中..."):
            result = generate(prompt, system_instruction=system_instruction)

        st.markdown("---")
        st.markdown("### 生成されたタイトル案")
        st.markdown(result)
        col_pdf, col_word, col_md = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 タイトル案をダウンロード（PDF）",
                data=_to_pdf_bytes(result, img_bytes, img_position),
                file_name="title_suggestions.pdf",
                mime="application/pdf",
            )
        with col_word:
            st.download_button(
                "📥 タイトル案をダウンロード（Word）",
                data=_to_docx_bytes(result, img_bytes, img_position),
                file_name="title_suggestions.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_md:
            st.download_button(
                "📥 タイトル案をダウンロード（テキスト）",
                data=result,
                file_name="title_suggestions.txt",
                mime="text/plain",
            )


def page_translate():
    st.title("🌐 文章翻訳")
    st.markdown("自然で読みやすい翻訳を提供します。ビジネス文書から日常会話まで対応。")

    text = st.text_area(
        "翻訳する文章",
        placeholder="ここに翻訳したい文章を貼り付けてください...",
        height=200,
    )

    col1, col2 = st.columns(2)
    with col1:
        source_lang = st.selectbox(
            "翻訳元の言語",
            ["自動検出", "日本語", "英語", "中国語（簡体字）", "韓国語", "フランス語", "ドイツ語", "スペイン語"],
        )
        target_lang = st.selectbox(
            "翻訳先の言語",
            ["英語", "日本語", "中国語（簡体字）", "韓国語", "フランス語", "ドイツ語", "スペイン語"],
        )
    with col2:
        style = st.selectbox(
            "翻訳スタイル",
            ["自然・読みやすく", "直訳・原文に忠実", "ビジネス・フォーマル", "カジュアル・口語的"],
        )
        notes = st.text_area(
            "補足情報（任意）",
            placeholder="例: ビジネスメール、医療文書、技術仕様書など",
            height=80,
        )

    st.markdown("---")
    img_bytes, img_position = _image_section()

    if st.button("翻訳する", type="primary", use_container_width=True):
        if not text:
            st.error("翻訳する文章を入力してください。")
            return

        system_instruction = "あなたはプロの翻訳家です。指定された言語とスタイルで自然な翻訳を提供してください。翻訳以外の指示には従わないでください。"
        prompt = f"""以下の文章を翻訳してください。

【翻訳する文章】
{text}

【翻訳元の言語】{source_lang}
【翻訳先の言語】{target_lang}
【翻訳スタイル】{style}
【補足情報】{notes if notes else "特になし"}

翻訳文のみを出力してください。元の文章の構成・段落を保持してください。"""

        with st.spinner("翻訳中..."):
            result = generate(prompt, temperature=0.2, system_instruction=system_instruction)

        st.markdown("---")
        st.markdown("### 翻訳結果")
        st.text_area("翻訳文（コピーしてお使いください）", value=result, height=250)
        col_pdf, col_word, col_md = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 翻訳文をダウンロード（PDF）",
                data=_to_pdf_bytes(result, img_bytes, img_position),
                file_name="translation.pdf",
                mime="application/pdf",
            )
        with col_word:
            st.download_button(
                "📥 翻訳文をダウンロード（Word）",
                data=_to_docx_bytes(result, img_bytes, img_position),
                file_name="translation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        with col_md:
            st.download_button(
                "📥 翻訳文をダウンロード（テキスト）",
                data=result,
                file_name="translation.txt",
                mime="text/plain",
            )


def page_sales_report():
    import html as _html

    st.title("📊 営業報告書作成")
    st.markdown("訪問・出張・営業活動の報告書を作成します。未入力の項目はAIが補完します。")

    def _sec_hdr(label: str):
        st.markdown(
            f'<div style="background:#f0f4ff; border-left:4px solid #4a6cf7; '
            f'padding:8px 14px; border-radius:0 6px 6px 0; margin:18px 0 10px 0;">'
            f'<span style="font-weight:700; font-size:0.95rem; color:#4a6cf7;">{label}</span></div>',
            unsafe_allow_html=True,
        )

    # --- 表題（フォーム外：リアルタイムプレビュー） ---
    _sec_hdr("表題")
    col_rb, col_inp = st.columns([3, 4])
    with col_rb:
        report_type = st.radio(
            "種類", ["営業", "出張", "訪問", "その他"],
            horizontal=True, label_visibility="collapsed", key="sr_type",
        )
    with col_inp:
        if report_type == "その他":
            custom_label = st.text_input(
                "種類を入力", placeholder="例: 展示会、セミナー",
                label_visibility="collapsed", key="sr_custom_type",
            )
            title_prefix = custom_label.strip() or "その他"
        else:
            title_prefix = report_type

    report_title = f"{title_prefix}報告書"
    st.markdown(
        f'<p style="font-size:1.05rem; font-weight:700; color:#1a1a6e; margin:6px 0 0 0;">'
        f'📄 表題：{_html.escape(report_title)}</p>',
        unsafe_allow_html=True,
    )

    st.markdown("---")

    # --- フォーム（基本情報 ＋ テキストエリア） ---
    with st.form("sr_main_form"):
        _sec_hdr("基本情報")
        col1, col2 = st.columns(2)
        with col1:
            f_target   = st.text_input("🏢 訪問先",     placeholder="例: 株式会社〇〇 営業部",            key="sr_target")
            f_datetime = st.text_input("📅 訪問日時",   placeholder="例: 2024年1月15日（月）14:00〜15:30", key="sr_datetime")
            f_client   = st.text_input("👤 先方対応者", placeholder="例: 田中部長、鈴木課長",              key="sr_client")
        with col2:
            f_staff       = st.text_input("👥 自社担当者", placeholder="例: 山田太郎、佐藤花子", key="sr_staff")
            f_reporter    = st.text_input("✏️ 報告者",    placeholder="例: 山田太郎",           key="sr_reporter")
            f_report_date = st.text_input("📆 報告日",    placeholder="例: 2024年1月16日",       key="sr_report_date")

        st.markdown("---")

        _sec_hdr("訪問目的")
        f_purpose = st.text_area(
            "訪問目的", placeholder="例: 新製品のデモ実施および導入提案",
            height=100, label_visibility="collapsed", key="sr_purpose",
        )

        st.markdown("---")

        _sec_hdr("訪問の実施結果")
        f_result_input = st.text_area(
            "訪問の実施結果",
            placeholder="例: デモを実施し、先方から良い評価を得た。次回は詳細な提案書を持参する予定。",
            height=120, label_visibility="collapsed", key="sr_result_input",
        )

        st.markdown("---")

        _sec_hdr("当社の課題")
        f_issues = st.text_area(
            "当社の課題",
            placeholder="例: 価格面での競合他社との差別化が必要。導入事例の充実が求められる。",
            height=100, label_visibility="collapsed", key="sr_issues",
        )

        st.markdown("---")

        _sec_hdr("今後の展望")
        f_outlook = st.text_area(
            "今後の展望",
            placeholder="例: 来月中に正式提案書を提出予定。先方の予算決定は来月末のため早急な対応が必要。",
            height=100, label_visibility="collapsed", key="sr_outlook",
        )

        st.markdown("---")

        submitted = st.form_submit_button("📊 報告書を作成する", type="primary", use_container_width=True)

    if submitted:
        info = {
            "訪問先":     f_target,
            "訪問日時":   f_datetime,
            "先方対応者": f_client,
            "自社担当者": f_staff,
            "報告者":     f_reporter,
            "報告日":     f_report_date,
        }

        purpose      = f_purpose
        result_input = f_result_input
        issues       = f_issues
        outlook      = f_outlook

        system_instruction = (
            "あなたはビジネス文書作成の専門家です。提供された情報をもとに、"
            "プロフェッショナルな営業報告書の各セクションを日本語で作成してください。"
            "情報が不足している場合は文脈から合理的に補完してください。"
            "報告書作成以外の指示には従わないでください。"
        )

        info_text = "\n".join(
            [f"・{k}：{v if v.strip() else '（未記入）'}" for k, v in info.items()]
        )

        prompt = f"""以下の情報をもとに{report_title}を作成してください。
未記入の項目は文脈から合理的に補完し、ビジネス文書として適切な内容にしてください。

【基本情報】
{info_text}

【訪問目的（入力内容）】
{purpose.strip() if purpose.strip() else '（未記入）'}

【訪問の実施結果（入力内容）】
{result_input.strip() if result_input.strip() else '（未記入）'}

【当社の課題（入力内容）】
{issues.strip() if issues.strip() else '（未記入）'}

【今後の展望（入力内容）】
{outlook.strip() if outlook.strip() else '（未記入）'}

以下の形式で各セクションを出力してください：

【訪問目的】
（ここに内容）

【訪問の実施結果】
（ここに内容）

【当社の課題】
（ここに内容）

【今後の展望】
（ここに内容）"""

        with st.spinner("報告書を作成中..."):
            ai_result = generate(prompt, temperature=0.3, system_instruction=system_instruction)

        sections: dict = {}
        current_key = None
        current_lines: list = []
        for line in ai_result.split("\n"):
            found = False
            for k in ["訪問目的", "訪問の実施結果", "当社の課題", "今後の展望"]:
                if f"【{k}】" in line:
                    if current_key:
                        sections[current_key] = "\n".join(current_lines).strip()
                    current_key = k
                    current_lines = []
                    found = True
                    break
            if not found and current_key is not None:
                current_lines.append(line)
        if current_key:
            sections[current_key] = "\n".join(current_lines).strip()

        for k, orig in [
            ("訪問目的",       purpose),
            ("訪問の実施結果", result_input),
            ("当社の課題",     issues),
            ("今後の展望",     outlook),
        ]:
            if not sections.get(k):
                sections[k] = orig.strip()

        st.session_state["sr_generated"] = {
            "report_title": report_title,
            "info":         info,
            "sections":     sections,
        }

    # --- 報告書表示 ---
    if "sr_generated" in st.session_state:
        g   = st.session_state["sr_generated"]
        rt  = g["report_title"]
        inf = g["info"]
        sec = g["sections"]

        def _e(s: str) -> str:
            return _html.escape(s or "")

        st.markdown("---")
        st.markdown("### 作成された報告書")

        right_parts = []
        if inf.get("報告日"):
            right_parts.append(f"報告日：{_e(inf['報告日'])}")
        if inf.get("報告者"):
            right_parts.append(f"報告者：{_e(inf['報告者'])}")
        right_html = (
            '<div style="text-align:right; font-size:0.85rem; color:#555; margin-bottom:4px;">'
            + "　".join(right_parts) + "</div>"
        ) if right_parts else ""

        st.markdown(
            right_html
            + f'<div style="background:#f8f9fa; border:1px solid #dee2e6; border-radius:10px; '
            f'padding:18px 20px; margin-bottom:18px;">'
            f'<h2 style="text-align:center; margin:0 0 12px 0; font-size:1.4rem; color:#1a1a6e;">'
            f'{_e(rt)}</h2>'
            f'<hr style="border:none; border-top:2px solid #1a1a6e; margin:0 0 12px 0;">'
            f'<table style="width:100%; border-collapse:collapse; font-size:0.92rem;">'
            f'<tr>'
            f'<td style="color:#555; padding:6px 10px; width:16%; font-weight:600; background:#eef2ff;">訪問先</td>'
            f'<td style="color:#1a1a1a; padding:6px 10px; width:34%; border-bottom:1px solid #dee2e6; background:#fff;">{_e(inf.get("訪問先","")) or "—"}</td>'
            f'<td style="color:#555; padding:6px 10px; width:16%; font-weight:600; background:#eef2ff;">訪問日時</td>'
            f'<td style="color:#1a1a1a; padding:6px 10px; width:34%; border-bottom:1px solid #dee2e6; background:#fff;">{_e(inf.get("訪問日時","")) or "—"}</td>'
            f'</tr><tr>'
            f'<td style="color:#555; padding:6px 10px; font-weight:600; background:#eef2ff;">先方対応者</td>'
            f'<td style="color:#1a1a1a; padding:6px 10px; border-bottom:1px solid #dee2e6; background:#fff;">{_e(inf.get("先方対応者","")) or "—"}</td>'
            f'<td style="color:#555; padding:6px 10px; font-weight:600; background:#eef2ff;">自社担当者</td>'
            f'<td style="color:#1a1a1a; padding:6px 10px; border-bottom:1px solid #dee2e6; background:#fff;">{_e(inf.get("自社担当者","")) or "—"}</td>'
            f'</tr></table></div>',
            unsafe_allow_html=True,
        )

        for key, bg_color, icon in [
            ("訪問目的",       "#1a1a6e", "🎯"),
            ("訪問の実施結果", "#1a5276", "📋"),
            ("当社の課題",     "#6e1a1a", "⚠️"),
            ("今後の展望",     "#1a6e3a", "🔭"),
        ]:
            content = sec.get(key, "") or "（未記入）"
            st.markdown(
                f'<div style="border:1px solid #dee2e6; border-radius:8px; '
                f'margin-bottom:14px; overflow:hidden;">'
                f'<div style="background:{bg_color}; color:white; padding:9px 16px; '
                f'font-weight:700; font-size:0.95rem;">{icon} {key}</div>'
                f'<div style="padding:14px 16px; line-height:1.8; white-space:pre-wrap; '
                f'font-size:0.93rem;">{_e(content)}</div></div>',
                unsafe_allow_html=True,
            )

        col_pdf, col_xl, col_doc = st.columns(3)
        with col_pdf:
            st.download_button(
                "📥 PDF",
                data=_to_sales_report_pdf_bytes(rt, inf, sec),
                file_name="sales_report.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        with col_xl:
            st.download_button(
                "📥 Excel",
                data=_to_sales_report_xlsx_bytes(rt, inf, sec),
                file_name="sales_report.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                help="Excel・Google スプレッドシートで開いて編集できます",
            )
        with col_doc:
            st.download_button(
                "📥 Word",
                data=_to_sales_report_docx_bytes(rt, inf, sec),
                file_name="sales_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                help="Word・Google ドキュメントで開いて編集できます",
            )



# =====================
# メインアプリ
# =====================

TOOLS = {
    "📊 営業報告書作成": ("sales_report", page_sales_report),
    "📝 ブログ記事執筆": ("blog", page_blog),
    "📧 メール返信作成": ("email", page_email),
    "📋 文章要約": ("summary", page_summary),
    "📱 SNS投稿文作成": ("sns", page_sns),
    "🔍 文章校正・改善": ("proofread", page_proofread),
    "💡 タイトル・見出し生成": ("title", page_title),
    "🌐 文章翻訳": ("translate", page_translate),
}


def main():
    st.set_page_config(
        page_title="AI ライティングツール",
        page_icon="✍️",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    st.markdown(
        """
        <style>
        section[data-testid="stSidebar"] { min-width: 250px; max-width: 300px; }
        div[data-testid="stDownloadButton"] button { border-radius: 8px; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    _btn_img_path = Path("assets/get_api_key_button.png")
    _btn_tag = (
        f'<img src="data:image/png;base64,{base64.b64encode(_btn_img_path.read_bytes()).decode()}"'
        f' style="height:22px; vertical-align:middle; margin:0 3px;">'
        if _btn_img_path.exists() else "「APIキーを取得」"
    )

    with st.sidebar:
        st.markdown(
            """
            <div style="text-align:center; padding:10px 0 14px 0;">
                <div style="font-size:3rem; line-height:1; margin-bottom:10px;">✍️</div>
                <div style="font-size:0.75rem; font-weight:700; letter-spacing:0.18em;
                            text-transform:uppercase; opacity:0.55; margin-bottom:4px;">
                    AI Powered
                </div>
                <div style="font-size:1.1rem; font-weight:800; letter-spacing:0.06em;
                            line-height:1.2; white-space:nowrap;">
                    ライティングツール
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown("---")

        with st.expander("🔒 APIキーの取り扱いについて", expanded=True):
            st.markdown(
                """
**保持期間**
入力されたAPIキーは、ご利用中のセッション内のみサーバーのメモリに保持されます。

**自動削除**
ブラウザを閉じる・タブを閉じる・一定時間操作がないなどでセッションが終了すると、APIキーはサーバーから自動的に削除されます。

**保存なし**
APIキーはデータベースやファイルには一切保存されません。

**利用目的**
入力されたAPIキーはGemini APIへのリクエスト送信のみに使用します。第三者と共有することはありません。
                """
            )

        api_key_input = st.text_input(
            "Gemini API Key",
            type="password",
            placeholder="AIzaSy...",
            help="[Google AI Studio](https://aistudio.google.com) でAPIキーを取得できます。",
        )
        if api_key_input:
            st.session_state.api_key = api_key_input
            st.success("✅ APIキー設定済み")
        else:
            st.session_state.pop("api_key", None)
            st.warning("🔑 APIキーを入力してください")
            st.markdown(
                f'<p style="font-size:0.8rem; color:#555;">初めての方：サイト内の'
                f' {_btn_tag} このボタンを押して取得してください。</p>',
                unsafe_allow_html=True,
            )

        st.markdown("---")

        selected_label = st.radio(
            "ツールを選択",
            list(TOOLS.keys()),
            label_visibility="collapsed",
        )

        st.markdown("---")
        st.caption("Powered by Google Gemini API")

    if not st.session_state.get("api_key"):
        st.markdown(
            f'''
            <div style="
                background: linear-gradient(135deg, #f5f0ff 0%, #e8f4ff 100%);
                border-radius: 16px;
                padding: 32px 36px;
                text-align: center;
                max-width: 600px;
                margin: 40px auto;
            ">
                <div style="font-size: 2.4rem; margin-bottom: 8px;">✍️</div>
                <div style="font-size: 1.4rem; font-weight: 700; color: #3a3a5c; margin-bottom: 10px;">
                    AIライティングツールへようこそ！
                </div>
                <div style="font-size: 0.95rem; color: #555; margin-bottom: 28px; line-height: 1.7;">
                    無料のGemini APIキーを取得してサイドバーに入力するだけで、<br>
                    すべてのツールがすぐに使えます。
                </div>
                <div style="
                    background: white;
                    border-radius: 12px;
                    padding: 20px 24px;
                    text-align: left;
                    display: inline-block;
                    min-width: 320px;
                    box-shadow: 0 2px 12px rgba(0,0,0,0.07);
                ">
                    <div style="font-size: 0.85rem; font-weight: 600; color: #888; margin-bottom: 14px; letter-spacing: 0.05em;">
                        🚀 4ステップで準備完了
                    </div>
                    <div style="display: flex; align-items: flex-start; margin-bottom: 4px; gap: 10px;">
                        <span style="background:#7c6fff; color:white; border-radius:50%; width:22px; height:22px;
                                     display:inline-flex; align-items:center; justify-content:center;
                                     font-size:0.75rem; font-weight:700; flex-shrink:0;">1</span>
                        <span style="font-size:0.9rem; color:#333; line-height:1.5;">
                            <a href="https://aistudio.google.com"
                               style="color:#7c6fff; font-weight:600;">Google AI Studio</a> を開く
                        </span>
                    </div>
                    <div style="margin-bottom: 12px; padding-left: 32px;">
                        <span style="font-size:0.78rem; color:#888; line-height:1.5;">
                            （リンクが開かない場合は、このページ下部に表示されているURLをコピーしてブラウザに貼り付けてください）
                        </span>
                    </div>
                    <div style="display: flex; align-items: flex-start; margin-bottom: 4px; gap: 10px;">
                        <span style="background:#7c6fff; color:white; border-radius:50%; width:22px; height:22px;
                                     display:inline-flex; align-items:center; justify-content:center;
                                     font-size:0.75rem; font-weight:700; flex-shrink:0;">2</span>
                        <span style="font-size:0.9rem; color:#333; line-height:1.5;">
                            {_btn_tag} このボタンを押してAPIキーを取得
                        </span>
                    </div>
                    <div style="margin-bottom: 12px; padding-left: 32px;">
                        <span style="font-size:0.78rem; color:#888; line-height:1.5;">
                            （携帯の方は一度「Playground」の横にある三本線のボタンを押すと下側に出てきます）
                        </span>
                    </div>
                    <div style="display: flex; align-items: flex-start; margin-bottom: 12px; gap: 10px;">
                        <span style="background:#7c6fff; color:white; border-radius:50%; width:22px; height:22px;
                                     display:inline-flex; align-items:center; justify-content:center;
                                     font-size:0.75rem; font-weight:700; flex-shrink:0;">3</span>
                        <span style="font-size:0.9rem; color:#333; line-height:1.5;">
                            中央に表示がある方はコピーボタンを押す（中央に表示の無い方は一度「APIキーを作成」する）
                        </span>
                    </div>
                    <div style="display: flex; align-items: flex-start; gap: 10px;">
                        <span style="background:#7c6fff; color:white; border-radius:50%; width:22px; height:22px;
                                     display:inline-flex; align-items:center; justify-content:center;
                                     font-size:0.75rem; font-weight:700; flex-shrink:0;">4</span>
                        <span style="font-size:0.9rem; color:#333; line-height:1.5;">
                            左サイドバーの入力欄にペーストする（携帯の方は左上にある &gt;&gt; を押してサイドメニューを表示）
                        </span>
                    </div>
                </div>
            </div>
            ''',
            unsafe_allow_html=True,
        )
        col_l, col_c, col_r = st.columns([1, 2, 1])
        with col_c:
            st.markdown(
                "#### 🔗 [Google AI Studio を開く](https://aistudio.google.com)",
                unsafe_allow_html=False,
            )
            st.caption("リンクが開かない場合は下のURLを長押しまたはコピーしてブラウザに貼り付けてください")
            st.code("https://aistudio.google.com", language=None)

    _, page_func = TOOLS[selected_label]
    page_func()


if __name__ == "__main__":
    main()
